#!/usr/bin/env python3
"""
ATS Resume Optimizer — Flask Web App
Wraps the CLI optimizer logic into a web UI deployable to GCP Cloud Run.
"""

import copy
import io
import json
import time
import os
import re
import tempfile
import requests as http_requests
from flask import Flask, request, render_template, send_file, jsonify, Response, stream_with_context
from docx import Document
from docx.oxml.ns import qn

app = Flask(__name__)

DEFAULT_API_KEY = os.environ.get("GEMINI_API_KEY", "AIzaSyDZBdPvjWxwq1dW5Mu3I9Bg5Xfi03XM034")
GEMINI_MODEL = "gemini-2.5-flash"
GCP_PROJECT = os.environ.get("GCP_PROJECT", "gen-lang-client-0682496991")
GCP_REGION = os.environ.get("GCP_REGION", "us-central1")
USE_VERTEX = os.environ.get("USE_VERTEX", "true").lower() == "true"
MIN_COUNT = 13
MAX_RETRIES = 5

# ═══════════════════════════════════════════════════════════
# Core optimizer logic (from ats_optimizer.py)
# ═══════════════════════════════════════════════════════════

def extract_sections(doc):
    """Extract resume sections from an already-loaded Document."""
    paragraphs = doc.paragraphs
    sections = {
        "summary": [], "client1_header": "", "client1_role": "",
        "client1_bullets": [], "client1_env": "",
        "client2_header": "", "client2_role": "",
        "client2_bullets": [], "client2_env": "",
        "other_sections_text": "",
    }

    summary_heading = summary_end = c1_start = c1_end = c2_start = c2_end = None

    for i, p in enumerate(paragraphs):
        txt = p.text.strip()
        if "PROFESSIONAL SUMMARY" in txt:
            summary_heading = i
        if summary_heading is not None and i > summary_heading and (
            txt.startswith("EDUCATION") or txt.startswith("Bachelor")
        ):
            if summary_end is None:
                summary_end = i
        if "Goodlabs" in txt and "Client" in txt:
            c1_start = i
        if c1_start and i > c1_start and "chevron" in txt.lower() and "Client" in txt:
            if c1_end is None:
                c1_end = i
        if "chevron" in txt.lower() and "Client" in txt:
            c2_start = i
        if c2_start and i > c2_start and "Client" in txt and ("JP Morgan" in txt or "Qualcomm" in txt):
            if c2_end is None:
                c2_end = i

    if summary_heading and summary_end:
        for i in range(summary_heading + 1, summary_end):
            txt = paragraphs[i].text.strip()
            if txt:
                sections["summary"].append(txt)

    if c1_start and c1_end:
        sections["client1_header"] = paragraphs[c1_start].text.strip()
        if c1_start + 1 < c1_end:
            sections["client1_role"] = paragraphs[c1_start + 1].text.strip()
        for i in range(c1_start + 2, c1_end):
            txt = paragraphs[i].text.strip()
            if txt:
                if txt.lower().startswith("environment"):
                    sections["client1_env"] = txt
                else:
                    sections["client1_bullets"].append(txt)

    if c2_start and c2_end:
        sections["client2_header"] = paragraphs[c2_start].text.strip()
        if c2_start + 1 < c2_end:
            sections["client2_role"] = paragraphs[c2_start + 1].text.strip()
        for i in range(c2_start + 2, c2_end):
            txt = paragraphs[i].text.strip()
            if txt:
                if txt.lower().startswith("environment"):
                    sections["client2_env"] = txt
                else:
                    sections["client2_bullets"].append(txt)

    if c2_end:
        other = []
        for i in range(c2_end, len(paragraphs)):
            txt = paragraphs[i].text.strip()
            if txt:
                other.append(txt)
        sections["other_sections_text"] = "\n".join(other)

    return sections


def _get_vertex_token():
    """Get access token from the GCP metadata server (works on Cloud Run)."""
    try:
        resp = http_requests.get(
            "http://metadata.google.internal/computeMetadata/v1/instance/service-accounts/default/token",
            headers={"Metadata-Flavor": "Google"},
            timeout=5,
        )
        resp.raise_for_status()
        return resp.json()["access_token"]
    except Exception:
        return None


def call_gemini(api_key, prompt):
    """Call Gemini — uses Vertex AI on Cloud Run, falls back to API key locally."""
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 0.7,
            "maxOutputTokens": 32768,
            "responseMimeType": "application/json",
        },
    }

    # Try Vertex AI first (higher rate limits, no API key needed on Cloud Run)
    token = None
    if USE_VERTEX:
        token = _get_vertex_token()

    if token:
        url = (
            f"https://{GCP_REGION}-aiplatform.googleapis.com/v1beta1/"
            f"projects/{GCP_PROJECT}/locations/{GCP_REGION}/"
            f"publishers/google/models/{GEMINI_MODEL}:generateContent"
        )
        headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}

        # Try Vertex AI, fall back to API key if it fails
        for retry in range(3):
            resp = http_requests.post(url, json=payload, headers=headers, timeout=120)
            if resp.status_code in (403, 429):
                wait = (retry + 1) * 5
                time.sleep(wait)
                continue
            if resp.status_code == 200:
                data = resp.json()
                return data["candidates"][0]["content"]["parts"][0]["text"]
            else:
                # Vertex AI failed — fall back to API key mode below
                break

    # Fallback: API key mode
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    for retry in range(3):
        resp = http_requests.post(url, json=payload, headers=headers, timeout=120)
        if resp.status_code in (403, 429):
            wait = (retry + 1) * 5
            time.sleep(wait)
            continue
        resp.raise_for_status()
        data = resp.json()
        return data["candidates"][0]["content"]["parts"][0]["text"]
    resp.raise_for_status()


def build_prompt(sections, keywords, feedback=None, job_description=None):
    kw_list = ", ".join(keywords)
    feedback_block = ""
    if feedback:
        lines = [f'  - "{kw}" currently appears {c} times (needs {MIN_COUNT}+)' for kw, c in feedback.items()]
        feedback_block = f"\n⚠️ CRITICAL: Fix these keyword counts:\n" + "\n".join(lines) + "\n"

    jd_block = ""
    if job_description:
        jd_block = f"""

=== TARGET JOB DESCRIPTION (tailor the resume to match this JD) ===
{job_description[:3000]}
"""

    return f"""You are an expert ATS resume optimizer for DevOps/SRE roles.

TASK: Rewrite Professional Summary, Client 1 (Goodlabs), and Client 2 (Chevron) for these keywords:
  {kw_list}

REQUIREMENTS:
1. Each keyword MUST appear at least {MIN_COUNT} times TOTAL across ALL sections combined.
2. Keep bullets professional and realistic.
3. Professional Summary: 12-16 bullet points.
4. Client 1 (Goodlabs): 16-20 bullet points + Environment line.
5. Client 2 (Chevron): 18-22 bullet points + Environment line.
6. Environment lines: comma-separated tech list including target keywords.
7. Preserve the candidate's actual experience context.
8. CRITICAL: Do NOT use any markdown formatting. No ** or * around words. Plain text only.
9. If a Job Description is provided below, tailor the bullets to closely match the JD's requirements and terminology.
{feedback_block}
CURRENT RESUME:

=== PROFESSIONAL SUMMARY ===
{chr(10).join("• " + b for b in sections["summary"])}

=== CLIENT 1: {sections["client1_header"]} ===
{sections["client1_role"]}
{chr(10).join("• " + b for b in sections["client1_bullets"])}
{sections["client1_env"]}

=== CLIENT 2: {sections["client2_header"]} ===
{sections["client2_role"]}
{chr(10).join("• " + b for b in sections["client2_bullets"])}
{sections["client2_env"]}

=== OTHER EXPERIENCE (context only) ===
{sections["other_sections_text"][:2000]}
{jd_block}

RESPOND WITH ONLY VALID JSON (no markdown fences):
{{
  "summary_bullets": ["..."],
  "client1_bullets": ["..."],
  "client1_env": "comma-separated list",
  "client2_bullets": ["..."],
  "client2_env": "comma-separated list"
}}
"""


def count_keywords(text, keywords):
    counts = {}
    text_lower = text.lower()
    for kw in keywords:
        counts[kw] = len(re.findall(re.escape(kw.lower()), text_lower))
    return counts


def all_text_from_content(content):
    parts = []
    parts.extend(content.get("summary_bullets", []))
    parts.extend(content.get("client1_bullets", []))
    parts.append(content.get("client1_env", ""))
    parts.extend(content.get("client2_bullets", []))
    parts.append(content.get("client2_env", ""))
    return " ".join(parts)


def strip_md(text):
    return text.replace("**", "").replace("*", "").strip()


def update_docx(doc, content, sections):
    """Update the Document in-place with new content."""
    paragraphs = doc.paragraphs

    summary_heading_idx = summary_end_idx = c1_start_idx = c1_end_idx = c2_start_idx = c2_end_idx = None

    for i, p in enumerate(paragraphs):
        txt = p.text.strip()
        if "PROFESSIONAL SUMMARY" in txt:
            summary_heading_idx = i
        if summary_heading_idx is not None and i > summary_heading_idx and (
            txt.startswith("EDUCATION") or txt.startswith("Bachelor")
        ):
            if summary_end_idx is None:
                summary_end_idx = i
        if "Goodlabs" in txt and "Client" in txt:
            c1_start_idx = i
        if c1_start_idx and i > c1_start_idx and "chevron" in txt.lower() and "Client" in txt:
            if c1_end_idx is None:
                c1_end_idx = i
        if "chevron" in txt.lower() and "Client" in txt:
            c2_start_idx = i
        if c2_start_idx and i > c2_start_idx and "Client" in txt and ("JP Morgan" in txt or "Qualcomm" in txt):
            if c2_end_idx is None:
                c2_end_idx = i

    summary_first_bullet = summary_heading_idx + 1
    while summary_first_bullet < summary_end_idx and not paragraphs[summary_first_bullet].text.strip():
        summary_first_bullet += 1

    bullet_template = None
    for i in range(summary_first_bullet, summary_end_idx):
        p = paragraphs[i]
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:numPr")) is not None:
            bullet_template = p
            break

    bold_template = paragraphs[c1_start_idx]

    def add_bullets_before(anchor_idx, bullets, env_line, header=None, role=None):
        anchor = paragraphs[anchor_idx]
        elements = []
        if header:
            elements.append(("bold", header))
        if role:
            elements.append(("bold", role))
        for b in bullets:
            elements.append(("bullet", b))
        if env_line:
            elements.append(("bold", "ENVIRONMENT:"))
            elements.append(("regular", env_line))
        elements.append(("empty", ""))

        for etype, text in elements:
            if etype == "bold":
                new_elem = copy.deepcopy(bold_template._element)
                for r in new_elem.findall(qn("w:r")):
                    new_elem.remove(r)
                anchor._element.addprevious(new_elem)
                from docx.text.paragraph import Paragraph
                new_para = Paragraph(new_elem, doc.element.body)
                run = new_para.add_run(text)
                run.bold = True
                run.font.name = "Cambria"
            elif etype == "bullet":
                new_elem = copy.deepcopy(bullet_template._element)
                for r in new_elem.findall(qn("w:r")):
                    new_elem.remove(r)
                anchor._element.addprevious(new_elem)
                from docx.text.paragraph import Paragraph
                new_para = Paragraph(new_elem, doc.element.body)
                run = new_para.add_run(text)
                run.bold = False
                run.font.name = "Cambria"
            elif etype == "regular":
                new_elem = copy.deepcopy(bold_template._element)
                for r in new_elem.findall(qn("w:r")):
                    new_elem.remove(r)
                anchor._element.addprevious(new_elem)
                from docx.text.paragraph import Paragraph
                new_para = Paragraph(new_elem, doc.element.body)
                run = new_para.add_run(text)
                run.bold = False
                run.font.name = "Cambria"
            elif etype == "empty":
                new_elem = copy.deepcopy(bold_template._element)
                for r in new_elem.findall(qn("w:r")):
                    new_elem.remove(r)
                anchor._element.addprevious(new_elem)

    # Client 2
    for i in range(c2_end_idx - 1, c2_start_idx - 1, -1):
        paragraphs[i]._element.getparent().remove(paragraphs[i]._element)
    paragraphs = doc.paragraphs
    c2_anchor = None
    for i, p in enumerate(paragraphs):
        if ("JP Morgan" in p.text or "Qualcomm" in p.text) and "Client" in p.text:
            c2_anchor = i
            break
    add_bullets_before(c2_anchor, content["client2_bullets"], content["client2_env"],
                       sections["client2_header"], sections["client2_role"])

    # Client 1
    paragraphs = doc.paragraphs
    c1_s = c1_e = None
    for i, p in enumerate(paragraphs):
        if "Goodlabs" in p.text and "Client" in p.text:
            c1_s = i
        if c1_s and i > c1_s and "chevron" in p.text.lower() and "Client" in p.text:
            c1_e = i
            break
    for i in range(c1_e - 1, c1_s - 1, -1):
        paragraphs[i]._element.getparent().remove(paragraphs[i]._element)
    paragraphs = doc.paragraphs
    c1_anchor = None
    for i, p in enumerate(paragraphs):
        if "chevron" in p.text.lower() and "Client" in p.text:
            c1_anchor = i
            break
    add_bullets_before(c1_anchor, content["client1_bullets"], content["client1_env"],
                       sections["client1_header"], sections["client1_role"])

    # Summary
    paragraphs = doc.paragraphs
    s_h = s_e = None
    for i, p in enumerate(paragraphs):
        if "PROFESSIONAL SUMMARY" in p.text:
            s_h = i
        if s_h is not None and i > s_h and (
            p.text.strip().startswith("EDUCATION") or p.text.strip().startswith("Bachelor")
        ):
            s_e = i
            break
    for i in range(s_e - 1, s_h, -1):
        paragraphs[i]._element.getparent().remove(paragraphs[i]._element)
    paragraphs = doc.paragraphs
    s_anchor = None
    for i, p in enumerate(paragraphs):
        if p.text.strip().startswith("EDUCATION") or p.text.strip().startswith("Bachelor"):
            s_anchor = i
            break
    anchor = paragraphs[s_anchor]
    for b in content["summary_bullets"]:
        new_elem = copy.deepcopy(bullet_template._element)
        for r in new_elem.findall(qn("w:r")):
            new_elem.remove(r)
        anchor._element.addprevious(new_elem)
        from docx.text.paragraph import Paragraph
        new_para = Paragraph(new_elem, doc.element.body)
        run = new_para.add_run(b)
        run.bold = False
        run.font.name = "Cambria"
    # empty line
    new_elem = copy.deepcopy(bold_template._element)
    for r in new_elem.findall(qn("w:r")):
        new_elem.remove(r)
    anchor._element.addprevious(new_elem)

    return doc


# In-memory store for generated files (file_id -> bytes)
import uuid
generated_files = {}


# ═══════════════════════════════════════════════════════════
# Flask Routes
# ═══════════════════════════════════════════════════════════

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/download/<file_id>/<filename>")
def download_file(file_id, filename):
    """Serve a generated .docx file by its ID."""
    if file_id not in generated_files:
        return jsonify({"error": "File not found or expired"}), 404
    file_bytes = generated_files.pop(file_id)  # One-time download, then clean up
    return send_file(
        io.BytesIO(file_bytes),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="Resume_ATS_Optimized.docx",
    )


@app.route("/optimize", methods=["POST"])
def optimize():
    """
    Accepts multipart form: resume (file), keywords (text), api_key (text).
    Returns SSE stream with progress, then a download link.
    """
    resume_file = request.files.get("resume")
    keywords_raw = request.form.get("keywords", "")
    api_key = request.form.get("api_key", DEFAULT_API_KEY).strip()
    job_description = request.form.get("job_description", "").strip()

    if not resume_file:
        return jsonify({"error": "No resume file uploaded"}), 400
    if not keywords_raw.strip():
        return jsonify({"error": "No keywords provided"}), 400

    keywords = [k.strip() for k in keywords_raw.split(",") if k.strip()]

    # Read file bytes into memory BEFORE the generator (file stream closes after request)
    file_bytes = resume_file.read()

    def generate():
        try:
            yield sse_msg("status", "📄 Reading resume...")
            doc = Document(io.BytesIO(file_bytes))
            sections = extract_sections(doc)
            yield sse_msg("info", f"Found {len(sections['summary'])} summary bullets, "
                          f"{len(sections['client1_bullets'])} Client 1 bullets, "
                          f"{len(sections['client2_bullets'])} Client 2 bullets")

            content = None
            feedback = None

            for attempt in range(1, MAX_RETRIES + 1):
                if attempt > 1:
                    yield sse_msg("info", "⏳ Waiting 3s before retry...")
                    time.sleep(3)
                yield sse_msg("status", f"🤖 Attempt {attempt}/{MAX_RETRIES} — Calling Gemini...")
                prompt = build_prompt(sections, keywords, feedback, job_description)

                try:
                    raw = call_gemini(api_key, prompt)
                except Exception as e:
                    yield sse_msg("error", f"Gemini API error: {str(e)}")
                    return

                raw = raw.strip()
                if raw.startswith("```"):
                    raw = re.sub(r"^```[a-z]*\n?", "", raw)
                    raw = re.sub(r"\n?```$", "", raw)

                try:
                    content = json.loads(raw)
                except json.JSONDecodeError:
                    # Try to extract JSON object from the response
                    match = re.search(r'\{[\s\S]*\}', raw)
                    if match:
                        try:
                            content = json.loads(match.group())
                        except json.JSONDecodeError:
                            yield sse_msg("warning", f"Attempt {attempt}: JSON parse error, retrying...")
                            feedback = {"JSON_ERROR": 0}
                            continue
                    else:
                        yield sse_msg("warning", f"Attempt {attempt}: JSON parse error, retrying...")
                        feedback = {"JSON_ERROR": 0}
                        continue

                required = ["summary_bullets", "client1_bullets", "client1_env", "client2_bullets", "client2_env"]
                if any(k not in content for k in required):
                    yield sse_msg("warning", f"Attempt {attempt}: Missing keys, retrying...")
                    feedback = {"MISSING_KEYS": 0}
                    continue

                # Strip markdown
                content["summary_bullets"] = [strip_md(b) for b in content["summary_bullets"]]
                content["client1_bullets"] = [strip_md(b) for b in content["client1_bullets"]]
                content["client2_bullets"] = [strip_md(b) for b in content["client2_bullets"]]
                content["client1_env"] = strip_md(content["client1_env"])
                content["client2_env"] = strip_md(content["client2_env"])

                all_text = all_text_from_content(content)
                counts = count_keywords(all_text, keywords)
                yield sse_msg("counts", json.dumps(counts))

                all_pass = all(c >= MIN_COUNT for c in counts.values())
                if all_pass:
                    yield sse_msg("status", "✅ All keywords meet the threshold!")
                    break
                else:
                    feedback = {kw: c for kw, c in counts.items() if c < MIN_COUNT}
                    failing = ", ".join(f"{kw}({c})" for kw, c in feedback.items())
                    yield sse_msg("warning", f"Below threshold: {failing} — retrying...")
            else:
                yield sse_msg("warning", "Max retries reached. Using best result.")

            if content is None:
                yield sse_msg("error", "Failed to generate valid content.")
                return

            yield sse_msg("status", "📝 Writing optimized resume...")

            fresh_doc = Document(io.BytesIO(file_bytes))
            updated_doc = update_docx(fresh_doc, content, sections)

            buf = io.BytesIO()
            updated_doc.save(buf)
            docx_bytes = buf.getvalue()

            # Store file and send download URL
            file_id = str(uuid.uuid4())
            generated_files[file_id] = docx_bytes
            yield sse_msg("download", f"/download/{file_id}/Resume_ATS_Optimized.docx")

            final_counts = count_keywords(all_text_from_content(content), keywords)
            yield sse_msg("final_counts", json.dumps(final_counts))
            yield sse_msg("done", "✅ Resume optimized successfully!")

        except Exception as e:
            yield sse_msg("error", f"Unexpected error: {str(e)}")

    return Response(stream_with_context(generate()), mimetype="text/event-stream")


def sse_msg(event, data):
    return f"event: {event}\ndata: {data}\n\n"


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port, debug=True)

